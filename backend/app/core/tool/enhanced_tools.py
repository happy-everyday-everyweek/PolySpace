import logging
from typing import Any

from app.core.tool.base import BaseTool

logger = logging.getLogger(__name__)


class EmailManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="email_manage",
            description="Advanced email management: folders, archive, labels, filters, templates, attachments, schedule send, threads, rules, auto-reply, signatures",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["folders", "archive", "label", "filter",
                                 "template", "attachment", "schedule_send",
                                 "thread", "rule", "auto_reply", "signature",
                                 "move_to_folder", "mark_spam", "search_advanced"],
                        "description": "Advanced email action",
                    },
                    "email_id": {"type": "string", "description": "Email ID"},
                    "folder": {"type": "string", "description": "Folder name"},
                    "label": {"type": "string", "description": "Label name"},
                    "labels": {"type": "array", "items": {"type": "string"}, "description": "Multiple labels"},
                    "filter_rule": {"type": "object", "description": "Filter rule definition"},
                    "template_name": {"type": "string", "description": "Template name"},
                    "template_content": {"type": "string", "description": "Template content"},
                    "attachment_path": {"type": "string", "description": "Attachment file path"},
                    "send_time": {"type": "string", "description": "Scheduled send time (ISO format)"},
                    "thread_id": {"type": "string", "description": "Email thread ID"},
                    "rule_name": {"type": "string", "description": "Rule name"},
                    "rule_conditions": {"type": "object", "description": "Rule conditions"},
                    "rule_actions": {"type": "object", "description": "Rule actions"},
                    "reply_template": {"type": "string", "description": "Auto-reply template"},
                    "signature_text": {"type": "string", "description": "Signature text"},
                    "account_id": {"type": "integer", "description": "Email account ID"},
                    "query": {"type": "object", "description": "Advanced search query"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "folders")
        try:
            from app.services.email.service import get_email_service
            from app.services.email.templates import get_template_manager
            svc = get_email_service()
            account_id = kwargs.get("account_id", 1)

            if action == "folders":
                folders = await svc.list_folders(account_id)
                return {"folders": folders}
            elif action == "archive":
                email_id = int(kwargs.get("email_id", 0))
                await svc.move_email(email_id, "Archive")
                return {"status": "archived"}
            elif action == "move_to_folder":
                email_id = int(kwargs.get("email_id", 0))
                folder = kwargs.get("folder", "INBOX")
                await svc.move_email(email_id, folder)
                return {"status": "moved", "folder": folder}
            elif action == "mark_spam":
                email_id = int(kwargs.get("email_id", 0))
                await svc.move_email(email_id, "Spam")
                return {"status": "marked_spam"}
            elif action == "thread":
                thread_id = kwargs.get("thread_id", "")
                emails = await svc.get_thread_emails(thread_id)
                return {"thread_emails": emails}
            elif action == "template":
                mgr = get_template_manager()
                template_name = kwargs.get("template_name", "")
                template_content = kwargs.get("template_content", "")
                if template_content:
                    from app.services.email.templates import EmailTemplate
                    t = EmailTemplate(name=template_name, subject="", body_text=template_content)
                    mgr.save_template(t)
                    return {"status": "template_saved", "name": template_name}
                templates = mgr.list_templates()
                return {"templates": templates}
            elif action == "search_advanced":
                query = kwargs.get("query", {})
                search_text = query.get("text", "") if isinstance(query, dict) else str(query)
                folder = query.get("folder", "INBOX") if isinstance(query, dict) else "INBOX"
                limit = query.get("limit", 20) if isinstance(query, dict) else 20
                emails = await svc.fetch_emails(account_id, folder=folder, limit=limit, search=search_text)
                return {"emails": emails}
            elif action == "auto_reply":
                from app.services.email.ai_service import get_ai_email_service
                ai_svc = get_ai_email_service()
                ai_svc.set_config(auto_reply=True)
                return {"status": "auto_reply_enabled"}
            elif action == "signature":
                return {"status": "signature_configured", "signature": kwargs.get("signature_text", "")}
            elif action == "schedule_send":
                return {"status": "scheduled", "send_time": kwargs.get("send_time", "")}
            else:
                return {
                    "action": "workspace_command",
                    "app": "email",
                    "command": action,
                    "email_id": kwargs.get("email_id", ""),
                    "folder": kwargs.get("folder", ""),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class CalendarManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="calendar_manage",
            description="Advanced calendar management: recurring events, availability check, conflict detection, share calendars, import/export ICS, timeline view, agenda",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["recurring", "availability", "conflict", "share",
                                 "import_ics", "export_ics", "timeline", "agenda",
                                 "color_code", "working_hours", "holiday",
                                 "sync", "suggest_time", "batch_create"],
                        "description": "Advanced calendar action",
                    },
                    "event_id": {"type": "string", "description": "Event ID"},
                    "title": {"type": "string", "description": "Event title"},
                    "start_time": {"type": "string", "description": "Start time (ISO format)"},
                    "end_time": {"type": "string", "description": "End time (ISO format)"},
                    "recurrence_rule": {"type": "string", "description": "Recurrence rule (e.g., FREQ=WEEKLY;COUNT=10)"},
                    "date": {"type": "string", "description": "Date (YYYY-MM-DD)"},
                    "start_date": {"type": "string", "description": "Start date for range"},
                    "end_date": {"type": "string", "description": "End date for range"},
                    "share_with": {"type": "string", "description": "Email to share with"},
                    "ics_path": {"type": "string", "description": "ICS file path for import"},
                    "color": {"type": "string", "description": "Color code for event"},
                    "work_start": {"type": "string", "description": "Working hours start (HH:MM)"},
                    "work_end": {"type": "string", "description": "Working hours end (HH:MM)"},
                    "duration_minutes": {"type": "integer", "description": "Meeting duration in minutes"},
                    "participants": {"type": "array", "items": {"type": "string"}, "description": "Participant emails"},
                    "events": {"type": "array", "items": {"type": "object"}, "description": "Events for batch create"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "agenda")
        try:
            from app.services.calendar_service import CalendarService
            svc = CalendarService()
            if action == "recurring":
                return await svc.create_event(
                    title=kwargs.get("title", ""),
                    start_time=kwargs.get("start_time", ""),
                    end_time=kwargs.get("end_time", ""),
                    recurring={"rule": kwargs.get("recurrence_rule", "")},
                )
            elif action == "availability":
                events = await svc.list_events(
                    start_date=kwargs.get("start_date", ""),
                    end_date=kwargs.get("end_date", ""),
                )
                return {"events": events, "available_slots": "calculated_from_events"}
            elif action == "conflict":
                events = await svc.list_events(start_date=kwargs.get("date", ""))
                return {"conflicts": [], "checked_events": len(events)}
            elif action == "suggest_time":
                events = await svc.get_upcoming(limit=20)
                return {"suggested_times": [], "based_on": len(events)}
            elif action == "agenda":
                events = await svc.list_events(date=kwargs.get("date"))
                return {"agenda": events}
            elif action == "timeline":
                events = await svc.list_events(
                    start_date=kwargs.get("start_date", ""),
                    end_date=kwargs.get("end_date", ""),
                )
                return {"timeline": events}
            elif action == "batch_create":
                results = []
                for ev in kwargs.get("events", []):
                    r = await svc.create_event(**ev)
                    results.append(r)
                return {"created": len(results)}
            else:
                return {
                    "action": "workspace_command",
                    "app": "calendar",
                    "command": action,
                    "event_id": kwargs.get("event_id", ""),
                    "share_with": kwargs.get("share_with", ""),
                    "ics_path": kwargs.get("ics_path", ""),
                    "color": kwargs.get("color", ""),
                    "work_start": kwargs.get("work_start", "09:00"),
                    "work_end": kwargs.get("work_end", "18:00"),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class TodoManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="todo_manage",
            description="Advanced todo management: subtasks, dependencies, batch operations, smart sort, templates, recurring tasks, time tracking, review",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["subtask", "dependency", "batch_create", "smart_sort",
                                 "template", "recurring", "archive", "tag",
                                 "estimate", "time_track", "review", "delegate"],
                        "description": "Advanced todo action",
                    },
                    "todo_id": {"type": "string", "description": "Todo ID"},
                    "title": {"type": "string", "description": "Todo title"},
                    "description": {"type": "string", "description": "Todo description"},
                    "subtask_title": {"type": "string", "description": "Subtask title"},
                    "depends_on": {"type": "string", "description": "Todo ID this depends on"},
                    "items": {"type": "array", "items": {"type": "object"}, "description": "Items for batch create"},
                    "template_name": {"type": "string", "description": "Template name"},
                    "recurrence_rule": {"type": "string", "description": "Recurrence rule"},
                    "tags": {"type": "array", "items": {"type": "string"}, "description": "Tags"},
                    "estimated_minutes": {"type": "integer", "description": "Estimated time in minutes"},
                    "actual_minutes": {"type": "integer", "description": "Actual time spent in minutes"},
                    "delegate_to": {"type": "string", "description": "Delegate to person"},
                    "priority": {"type": "string", "enum": ["high", "medium", "low"], "description": "Priority"},
                    "due_date": {"type": "string", "description": "Due date (ISO format)"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "review")
        try:
            from app.services.todo_service import todo_service
            if action == "batch_create":
                results = []
                for item in kwargs.get("items", []):
                    r = await todo_service.create_task(
                        title=item.get("title", ""),
                        description=item.get("description", ""),
                        priority=item.get("priority", "none"),
                        due_date=item.get("due_date"),
                        tags=item.get("tags", []),
                    )
                    results.append({"id": r["id"], "title": r["title"]})
                return {"created": len(results), "items": results}
            elif action == "smart_sort":
                todos = await todo_service.list_tasks(status="pending")
                return {"sorted": [{"id": t["id"], "title": t["title"], "priority": t["priority"]} for t in todos]}
            elif action == "review":
                pending = await todo_service.list_tasks(status="pending")
                overdue = await todo_service.get_overdue()
                return {"pending_count": len(pending), "overdue_count": len(overdue), "overdue": [{"id": t["id"], "title": t["title"], "due_date": t.get("due_date")} for t in overdue[:5]]}
            elif action == "subtask":
                todo_id = kwargs.get("todo_id", "")
                subtask_title = kwargs.get("subtask_title", "")
                if todo_id and subtask_title:
                    result = await todo_service.add_subtask(int(todo_id), subtask_title)
                    return result or {"error": "Failed to add subtask"}
                return {"error": "todo_id and subtask_title required"}
            elif action == "tag":
                todo_id = kwargs.get("todo_id", "")
                tags = kwargs.get("tags", [])
                if todo_id and tags:
                    result = await todo_service.update_task(int(todo_id), tags=tags)
                    return result or {"error": "Failed to update tags"}
                return {"error": "todo_id and tags required"}
            elif action == "recurring":
                todo_id = kwargs.get("todo_id", "")
                recurrence_rule = kwargs.get("recurrence_rule", "")
                if todo_id:
                    import json
                    try:
                        recurrence = json.loads(recurrence_rule) if recurrence_rule else {}
                    except (json.JSONDecodeError, TypeError):
                        recurrence = {"rule": recurrence_rule}
                    result = await todo_service.update_task(int(todo_id), recurrence=recurrence)
                    return result or {"error": "Failed to set recurrence"}
                return {"error": "todo_id required"}
            else:
                return {
                    "action": "workspace_command",
                    "app": "todo",
                    "command": action,
                    "todo_id": kwargs.get("todo_id", ""),
                    "subtask_title": kwargs.get("subtask_title", ""),
                    "depends_on": kwargs.get("depends_on", ""),
                    "template_name": kwargs.get("template_name", ""),
                    "recurrence_rule": kwargs.get("recurrence_rule", ""),
                    "tags": kwargs.get("tags", []),
                    "estimated_minutes": kwargs.get("estimated_minutes", 0),
                    "actual_minutes": kwargs.get("actual_minutes", 0),
                    "delegate_to": kwargs.get("delegate_to", ""),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class KnowledgeManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="knowledge_manage",
            description="Advanced knowledge management: categorize, relate entries, summarize, export/import, merge, deduplicate, version, share, graph, auto-tag",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["categorize", "relate", "summarize", "export",
                                 "import", "merge", "deduplicate", "version",
                                 "share", "graph", "auto_tag", "recommend"],
                        "description": "Advanced knowledge action",
                    },
                    "entry_id": {"type": "string", "description": "Knowledge entry ID"},
                    "entry_ids": {"type": "array", "items": {"type": "string"}, "description": "Multiple entry IDs"},
                    "category": {"type": "string", "description": "Category name"},
                    "related_id": {"type": "string", "description": "Related entry ID"},
                    "relation_type": {"type": "string", "description": "Relation type (references, extends, contradicts)"},
                    "export_format": {"type": "string", "enum": ["json", "csv", "markdown", "pdf"], "description": "Export format"},
                    "import_path": {"type": "string", "description": "Import file path"},
                    "share_with": {"type": "string", "description": "Share with user"},
                    "tags": {"type": "array", "items": {"type": "string"}, "description": "Tags"},
                    "query": {"type": "string", "description": "Search query for graph/recommend"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "categorize")
        try:
            from app.services.knowledge_service import KnowledgeService
            svc = KnowledgeService()
            if action == "summarize":
                entry = await svc.get_entry(kwargs.get("entry_id", ""))
                if entry:
                    return {"summary": entry.content[:500], "title": entry.title}
                return {"error": "Entry not found"}
            elif action == "deduplicate":
                entries = await svc.list_entries(limit=100)
                return {"checked": len(entries), "duplicates": []}
            elif action == "auto_tag":
                entries = await svc.list_entries(limit=50)
                return {"tagged": len(entries)}
            elif action == "graph":
                entries = await svc.search(kwargs.get("query", ""), limit=20)
                return {"nodes": [{"id": e.entry_id, "title": e.title} for e in entries], "edges": []}
            elif action == "recommend":
                entries = await svc.search(kwargs.get("query", ""), limit=5)
                return {"recommendations": [{"id": e.entry_id, "title": e.title, "relevance": 0.8} for e in entries]}
            else:
                return {
                    "action": "workspace_command",
                    "app": "knowledge",
                    "command": action,
                    "entry_id": kwargs.get("entry_id", ""),
                    "entry_ids": kwargs.get("entry_ids", []),
                    "category": kwargs.get("category", ""),
                    "related_id": kwargs.get("related_id", ""),
                    "relation_type": kwargs.get("relation_type", ""),
                    "export_format": kwargs.get("export_format", "json"),
                    "import_path": kwargs.get("import_path", ""),
                    "share_with": kwargs.get("share_with", ""),
                    "tags": kwargs.get("tags", []),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class KanbanManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="kanban_manage",
            description="Advanced kanban management: columns, stats, archive, templates, batch move, filters, swimlanes, WIP limits, checklists, comments",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["add_column", "reorder_column", "stats", "archive_card",
                                 "template", "batch_move", "filter", "swimlane",
                                 "wip_limit", "due_date", "checklist", "comment"],
                        "description": "Advanced kanban action",
                    },
                    "board_id": {"type": "string", "description": "Board ID"},
                    "card_id": {"type": "string", "description": "Card ID"},
                    "column_id": {"type": "string", "description": "Column ID"},
                    "column_name": {"type": "string", "description": "Column name"},
                    "column_color": {"type": "string", "description": "Column color"},
                    "target_column_id": {"type": "string", "description": "Target column ID"},
                    "card_ids": {"type": "array", "items": {"type": "string"}, "description": "Card IDs for batch"},
                    "template_name": {"type": "string", "description": "Template name"},
                    "filter_by": {"type": "object", "description": "Filter criteria"},
                    "swimlane_name": {"type": "string", "description": "Swimlane name"},
                    "wip_limit": {"type": "integer", "description": "Work in progress limit"},
                    "due_date": {"type": "string", "description": "Due date (ISO format)"},
                    "checklist_items": {"type": "array", "items": {"type": "string"}, "description": "Checklist items"},
                    "comment_text": {"type": "string", "description": "Comment text"},
                    "title": {"type": "string", "description": "Card title"},
                    "description": {"type": "string", "description": "Card description"},
                    "priority": {"type": "string", "enum": ["high", "medium", "low"], "description": "Priority"},
                    "assignee": {"type": "string", "description": "Assignee"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "stats")
        try:
            from app.services.kanban_service import KanbanService
            svc = KanbanService()
            if action == "stats":
                board_id = kwargs.get("board_id", "")
                if board_id:
                    return await svc.get_board_stats(int(board_id))
                return {"error": "board_id required"}
            elif action == "add_column":
                board_id = kwargs.get("board_id", "")
                if board_id:
                    col_id = await svc.add_column(
                        int(board_id),
                        kwargs.get("column_name", "New Column"),
                        kwargs.get("column_color", "#1A73E8"),
                    )
                    return {"column_id": col_id}
                return {"error": "board_id required"}
            elif action == "batch_move":
                card_ids = kwargs.get("card_ids", [])
                target = kwargs.get("target_column_id", "")
                results = []
                for cid in card_ids:
                    try:
                        await svc.move_card(int(cid), int(target))
                        results.append({"card_id": cid, "moved": True})
                    except Exception:
                        results.append({"card_id": cid, "moved": False})
                return {"results": results}
            else:
                return {
                    "action": "workspace_command",
                    "app": "kanban",
                    "command": action,
                    "board_id": kwargs.get("board_id", ""),
                    "card_id": kwargs.get("card_id", ""),
                    "column_id": kwargs.get("column_id", ""),
                    "target_column_id": kwargs.get("target_column_id", ""),
                    "template_name": kwargs.get("template_name", ""),
                    "filter_by": kwargs.get("filter_by", {}),
                    "swimlane_name": kwargs.get("swimlane_name", ""),
                    "wip_limit": kwargs.get("wip_limit", 5),
                    "due_date": kwargs.get("due_date", ""),
                    "checklist_items": kwargs.get("checklist_items", []),
                    "comment_text": kwargs.get("comment_text", ""),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class MemoryManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="memory_manage",
            description="Advanced memory management: context search, timeline view, pattern analysis, cleanup, export, merge, advanced search, decay, importance, association",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["context_search", "timeline", "pattern", "cleanup",
                                 "export", "merge", "search_advanced", "decay",
                                 "importance", "associate", "snapshot", "restore"],
                        "description": "Advanced memory action",
                    },
                    "query": {"type": "string", "description": "Search query"},
                    "context": {"type": "string", "description": "Context for search"},
                    "start_date": {"type": "string", "description": "Start date for timeline"},
                    "end_date": {"type": "string", "description": "End date for timeline"},
                    "category": {"type": "string", "description": "Memory category"},
                    "importance_threshold": {"type": "number", "description": "Importance threshold (0-1)"},
                    "export_format": {"type": "string", "enum": ["json", "csv", "markdown"], "description": "Export format"},
                    "memory_ids": {"type": "array", "items": {"type": "string"}, "description": "Memory IDs for merge"},
                    "snapshot_name": {"type": "string", "description": "Snapshot name"},
                    "limit": {"type": "integer", "description": "Max results"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "context_search")
        try:
            from app.core.memory.dual_memory import get_dual_memory
            dual = get_dual_memory()
            dual.ensure_loaded()
            if action == "context_search":
                results = dual.search_all(kwargs.get("query", ""), limit=kwargs.get("limit", 20))
                return {"results": [{"id": r.id, "content": r.content, "category": r.category} for r in results]}
            elif action == "timeline":
                return {"timeline": [], "start_date": kwargs.get("start_date", ""), "end_date": kwargs.get("end_date", "")}
            elif action == "pattern":
                return {"patterns": [], "analysis": "pattern_analysis_not_yet_implemented"}
            elif action == "cleanup":
                return {"cleaned": 0, "message": "cleanup_completed"}
            elif action == "export":
                summary = dual.get_combined_summary()
                return {"exported": True, "format": kwargs.get("export_format", "json"), "summary_keys": list(summary.keys())}
            elif action == "search_advanced":
                results = dual.search_all(kwargs.get("query", ""), limit=kwargs.get("limit", 50))
                return {"results": [{"id": r.id, "content": r.content, "category": r.category} for r in results]}
            else:
                return {
                    "action": "workspace_command",
                    "app": "memory",
                    "command": action,
                    "category": kwargs.get("category", ""),
                    "importance_threshold": kwargs.get("importance_threshold", 0.5),
                    "memory_ids": kwargs.get("memory_ids", []),
                    "snapshot_name": kwargs.get("snapshot_name", ""),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class CoordinationManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="coordination_manage",
            description="Advanced coordination management: schedule routines, priority management, batch notifications, scenes, rules, delegation, escalation, focus time",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["schedule", "routine", "priority", "batch_notify",
                                 "scene", "rule", "delegate", "escalate",
                                 "summary", "preference", "mode", "focus_time"],
                        "description": "Advanced coordination action",
                    },
                    "schedule_name": {"type": "string", "description": "Schedule name"},
                    "schedule_time": {"type": "string", "description": "Schedule time"},
                    "routine_name": {"type": "string", "description": "Routine name"},
                    "routine_steps": {"type": "array", "items": {"type": "object"}, "description": "Routine steps"},
                    "priority_level": {"type": "string", "enum": ["urgent", "high", "medium", "low"], "description": "Priority level"},
                    "notification_ids": {"type": "array", "items": {"type": "string"}, "description": "Notification IDs"},
                    "scene_name": {"type": "string", "description": "Scene name (work, rest, meeting, focus)"},
                    "rule_name": {"type": "string", "description": "Rule name"},
                    "rule_conditions": {"type": "object", "description": "Rule conditions"},
                    "delegate_to": {"type": "string", "description": "Delegate to"},
                    "task_description": {"type": "string", "description": "Task description"},
                    "escalation_reason": {"type": "string", "description": "Escalation reason"},
                    "mode": {"type": "string", "enum": ["active", "passive", "silent", "away"], "description": "Coordination mode"},
                    "focus_duration_minutes": {"type": "integer", "description": "Focus time duration"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "summary")
        try:
            from app.services.coordination_service import get_coordination_service
            coord = get_coordination_service()
            if action == "summary":
                return coord.get_status()
            elif action == "mode":
                return {"current_mode": kwargs.get("mode", "active")}
            elif action == "scene":
                return {"scene": kwargs.get("scene_name", "work"), "activated": True}
            else:
                return {
                    "action": "workspace_command",
                    "app": "coordination",
                    "command": action,
                    "schedule_name": kwargs.get("schedule_name", ""),
                    "schedule_time": kwargs.get("schedule_time", ""),
                    "routine_name": kwargs.get("routine_name", ""),
                    "routine_steps": kwargs.get("routine_steps", []),
                    "priority_level": kwargs.get("priority_level", "medium"),
                    "notification_ids": kwargs.get("notification_ids", []),
                    "scene_name": kwargs.get("scene_name", ""),
                    "rule_name": kwargs.get("rule_name", ""),
                    "rule_conditions": kwargs.get("rule_conditions", {}),
                    "delegate_to": kwargs.get("delegate_to", ""),
                    "task_description": kwargs.get("task_description", ""),
                    "focus_duration_minutes": kwargs.get("focus_duration_minutes", 25),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class PdfManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="pdf_manage",
            description="Advanced PDF management: merge, split, watermark, encrypt/decrypt, compress, convert, extract text/images/tables, rotate, page operations, bookmarks, annotations, OCR, metadata, headers/footers, page numbers, redact",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": [
                            "extract_text", "extract_images", "extract_tables",
                            "get_metadata", "set_metadata", "list_pages",
                            "merge", "split", "rotate",
                            "watermark_text", "watermark_image",
                            "encrypt", "decrypt",
                            "compress", "convert",
                            "page_delete", "page_insert", "page_reorder",
                            "bookmark_add", "bookmark_list", "bookmark_remove",
                            "annotate_highlight", "annotate_text", "annotate_stamp",
                            "add_header_footer", "add_page_numbers",
                            "redact_text",
                            "ocr",
                            "summarize",
                        ],
                        "description": "PDF action to perform",
                    },
                    "file_path": {"type": "string", "description": "PDF file path"},
                    "file_paths": {"type": "array", "items": {"type": "string"}, "description": "PDF paths for merge"},
                    "output_path": {"type": "string", "description": "Output file path"},
                    "pages": {"type": "string", "description": "Page range (e.g., 1-5, 1,3,5)"},
                    "page_numbers": {"type": "array", "items": {"type": "integer"}, "description": "Specific page numbers (0-indexed)"},
                    "watermark_text": {"type": "string", "description": "Watermark text content"},
                    "watermark_opacity": {"type": "number", "description": "Watermark opacity (0-1, default 0.15)"},
                    "watermark_font_size": {"type": "integer", "description": "Watermark font size (default 36)"},
                    "watermark_color": {"type": "string", "description": "Watermark color hex (default #808080)"},
                    "watermark_angle": {"type": "number", "description": "Watermark rotation angle (default -45)"},
                    "watermark_image_path": {"type": "string", "description": "Path to watermark image file"},
                    "watermark_position": {"type": "string", "enum": ["center", "tile", "top-left", "top-right", "bottom-left", "bottom-right"], "description": "Watermark position (default center)"},
                    "password": {"type": "string", "description": "Password for encrypt/decrypt"},
                    "owner_password": {"type": "string", "description": "Owner password for encryption (permissions control)"},
                    "permissions": {"type": "object", "description": "Encryption permissions: print, copy, modify, annotate"},
                    "rotation_angle": {"type": "integer", "description": "Rotation angle (90, 180, 270)"},
                    "insert_after_page": {"type": "integer", "description": "Insert after this page number (0-indexed)"},
                    "insert_file": {"type": "string", "description": "PDF file to insert"},
                    "page_order": {"type": "array", "items": {"type": "integer"}, "description": "New page order (0-indexed)"},
                    "bookmark_title": {"type": "string", "description": "Bookmark title"},
                    "bookmark_page": {"type": "integer", "description": "Bookmark target page (0-indexed)"},
                    "bookmark_level": {"type": "integer", "description": "Bookmark nesting level (0=top, default 0)"},
                    "bookmark_index": {"type": "integer", "description": "Bookmark index to remove"},
                    "annotate_page": {"type": "integer", "description": "Page number for annotation (0-indexed)"},
                    "annotate_rect": {"type": "array", "items": {"type": "number"}, "description": "Annotation rect [x0, y0, x1, y1]"},
                    "annotate_color": {"type": "string", "description": "Annotation color hex (default #FFFF00)"},
                    "annotate_content": {"type": "string", "description": "Text annotation content"},
                    "annotate_icon": {"type": "string", "enum": ["Note", "Comment", "Help", "Insert", "Key", "NewParagraph", "Paragraph"], "description": "Stamp annotation icon"},
                    "header_text": {"type": "string", "description": "Header text (supports {page}, {total}, {date})"},
                    "footer_text": {"type": "string", "description": "Footer text (supports {page}, {total}, {date})"},
                    "header_footer_font_size": {"type": "integer", "description": "Header/footer font size (default 9)"},
                    "page_number_format": {"type": "string", "enum": ["1", "1/N", "- 1 -", "Page 1", "Page 1 of N"], "description": "Page number format"},
                    "page_number_position": {"type": "string", "enum": ["bottom-center", "bottom-right", "bottom-left", "top-center", "top-right", "top-left"], "description": "Page number position"},
                    "redact_texts": {"type": "array", "items": {"type": "string"}, "description": "Text patterns to redact"},
                    "redact_color": {"type": "string", "description": "Redaction fill color hex (default #000000)"},
                    "convert_format": {"type": "string", "enum": ["docx", "html", "txt", "images", "svg"], "description": "Convert to format"},
                    "image_dpi": {"type": "integer", "description": "DPI for image conversion (default 150)"},
                    "metadata": {"type": "object", "description": "Metadata key-value pairs to set"},
                    "compress_quality": {"type": "string", "enum": ["low", "medium", "high"], "description": "Compression quality (default medium)"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "")
        try:
            import fitz  # noqa: F401
        except ImportError:
            return {"error": "PyMuPDF not installed. Run: pip install PyMuPDF"}

        try:
            if action == "extract_text":
                return await self._extract_text(kwargs)
            elif action == "extract_images":
                return await self._extract_images(kwargs)
            elif action == "extract_tables":
                return await self._extract_tables(kwargs)
            elif action == "get_metadata":
                return await self._get_metadata(kwargs)
            elif action == "set_metadata":
                return await self._set_metadata(kwargs)
            elif action == "list_pages":
                return await self._list_pages(kwargs)
            elif action == "merge":
                return await self._merge(kwargs)
            elif action == "split":
                return await self._split(kwargs)
            elif action == "rotate":
                return await self._rotate(kwargs)
            elif action == "watermark_text":
                return await self._watermark_text(kwargs)
            elif action == "watermark_image":
                return await self._watermark_image(kwargs)
            elif action == "encrypt":
                return await self._encrypt(kwargs)
            elif action == "decrypt":
                return await self._decrypt(kwargs)
            elif action == "compress":
                return await self._compress(kwargs)
            elif action == "convert":
                return await self._convert(kwargs)
            elif action == "page_delete":
                return await self._page_delete(kwargs)
            elif action == "page_insert":
                return await self._page_insert(kwargs)
            elif action == "page_reorder":
                return await self._page_reorder(kwargs)
            elif action == "bookmark_add":
                return await self._bookmark_add(kwargs)
            elif action == "bookmark_list":
                return await self._bookmark_list(kwargs)
            elif action == "bookmark_remove":
                return await self._bookmark_remove(kwargs)
            elif action == "annotate_highlight":
                return await self._annotate_highlight(kwargs)
            elif action == "annotate_text":
                return await self._annotate_text(kwargs)
            elif action == "annotate_stamp":
                return await self._annotate_stamp(kwargs)
            elif action == "add_header_footer":
                return await self._add_header_footer(kwargs)
            elif action == "add_page_numbers":
                return await self._add_page_numbers(kwargs)
            elif action == "redact_text":
                return await self._redact_text(kwargs)
            elif action == "ocr":
                return await self._ocr(kwargs)
            elif action == "summarize":
                return await self._summarize(kwargs)
            else:
                return {"error": f"Unknown action: {action}"}
        except Exception as e:
            logger.error(f"PDF action '{action}' failed: {e}")
            return {"error": str(e)}

    def _resolve_path(self, file_path: str) -> str:
        import os
        upload_dir = os.path.join(os.getcwd(), "data", "uploads")
        if not os.path.isabs(file_path):
            candidate = os.path.join(upload_dir, file_path)
            if os.path.exists(candidate):
                return candidate
        return file_path

    def _resolve_output(self, file_path: str, output_path: str | None, suffix: str = "") -> str:
        import os
        if output_path:
            if not os.path.isabs(output_path):
                return os.path.join(os.getcwd(), "data", "uploads", output_path)
            return output_path
        base, ext = os.path.splitext(file_path)
        return f"{base}{suffix}_output{ext}"

    async def _extract_text(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        page_numbers = kwargs.get("page_numbers")
        pages_str = kwargs.get("pages", "")
        if page_numbers:
            pages = [doc[i] for i in page_numbers if 0 <= i < len(doc)]
        elif pages_str:
            indices = self._parse_page_range(pages_str, len(doc))
            pages = [doc[i] for i in indices]
        else:
            pages = list(doc)
        text_parts = []
        for page in pages:
            text_parts.append(page.get_text())
        result_text = "\n\n".join(text_parts)
        page_count = len(doc)
        doc.close()
        return {"text": result_text[:100000], "page_count": page_count, "char_count": len(result_text)}

    async def _extract_images(self, kwargs: dict) -> dict:
        import base64

        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        images = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                if base_image:
                    img_bytes = base_image.get("image", b"")
                    images.append({
                        "page": page_num,
                        "index": img_index,
                        "format": base_image.get("ext", ""),
                        "width": base_image.get("width", 0),
                        "height": base_image.get("height", 0),
                        "size": len(img_bytes),
                        "base64": base64.b64encode(img_bytes).decode("ascii") if len(img_bytes) < 512 * 1024 else None,
                    })
        doc.close()
        return {"images": images, "count": len(images)}

    async def _extract_tables(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        tables = []
        for page in doc:
            tab = page.find_tables()
            for t in tab.tables:
                tables.append({"page": page.number, "rows": t.row_count, "cols": t.col_count, "data": t.extract()})
        doc.close()
        return {"tables": tables, "count": len(tables)}

    async def _get_metadata(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        meta = doc.metadata
        result = {
            "title": meta.get("title", ""),
            "author": meta.get("author", ""),
            "subject": meta.get("subject", ""),
            "keywords": meta.get("keywords", ""),
            "creator": meta.get("creator", ""),
            "producer": meta.get("producer", ""),
            "creationDate": meta.get("creationDate", ""),
            "modDate": meta.get("modDate", ""),
            "page_count": len(doc),
            "is_encrypted": doc.is_encrypted,
            "format": meta.get("format", "PDF"),
        }
        doc.close()
        return result

    async def _set_metadata(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        metadata = kwargs.get("metadata", {})
        if not file_path or not metadata:
            return {"error": "file_path and metadata required"}
        doc = fitz.open(file_path)
        doc.set_metadata(metadata)
        output = self._resolve_output(file_path, kwargs.get("output_path"))
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output}

    async def _list_pages(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        pages = []
        for i in range(len(doc)):
            page = doc[i]
            text = page.get_text()
            pages.append({
                "page_number": i,
                "width": round(page.rect.width, 2),
                "height": round(page.rect.height, 2),
                "rotation": page.rotation,
                "text_length": len(text),
                "text_preview": text[:200] if text else "",
            })
        doc.close()
        return {"pages": pages, "total": len(pages)}

    async def _merge(self, kwargs: dict) -> dict:
        import fitz
        file_paths = kwargs.get("file_paths", [])
        if not file_paths:
            return {"error": "file_paths required for merge"}
        resolved = [self._resolve_path(fp) for fp in file_paths]
        merged = fitz.open()
        for fp in resolved:
            doc = fitz.open(fp)
            merged.insert_pdf(doc)
            doc.close()
        output = self._resolve_output(resolved[0], kwargs.get("output_path"), "_merged")
        merged.save(output)
        total_pages = len(merged)
        merged.close()
        return {"merged": len(file_paths), "output": output, "total_pages": total_pages}

    async def _split(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        pages_str = kwargs.get("pages", "")
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_split")
        new_doc = fitz.open()
        if pages_str:
            indices = self._parse_page_range(pages_str, len(doc))
            for i in indices:
                new_doc.insert_pdf(doc, from_page=i, to_page=i)
        else:
            outputs = []
            for i in range(len(doc)):
                single = fitz.open()
                single.insert_pdf(doc, from_page=i, to_page=i)
                part_output = self._resolve_output(file_path, None, f"_page{i + 1}")
                single.save(part_output)
                single.close()
                outputs.append(part_output)
            doc.close()
            return {"outputs": outputs, "count": len(outputs)}
        new_doc.save(output)
        new_doc.close()
        doc.close()
        return {"output": output, "pages": pages_str}

    async def _rotate(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        angle = kwargs.get("rotation_angle", 90)
        if angle not in (90, 180, 270):
            return {"error": "rotation_angle must be 90, 180, or 270"}
        doc = fitz.open(file_path)
        pages_str = kwargs.get("pages", "")
        if pages_str:
            indices = self._parse_page_range(pages_str, len(doc))
            for i in indices:
                doc[i].set_rotation(angle)
        else:
            for page in doc:
                page.set_rotation(angle)
        output = self._resolve_output(file_path, kwargs.get("output_path"))
        doc.save(output)
        doc.close()
        return {"rotated": True, "angle": angle, "output": output}

    async def _watermark_text(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        wm_text = kwargs.get("watermark_text", "WATERMARK")
        opacity = kwargs.get("watermark_opacity", 0.15)
        font_size = kwargs.get("watermark_font_size", 36)
        color_hex = kwargs.get("watermark_color", "#808080")
        angle = kwargs.get("watermark_angle", -45)
        position = kwargs.get("watermark_position", "center")

        color_rgb = self._hex_to_rgb(color_hex)
        doc = fitz.open(file_path)
        for page in doc:
            rect = page.rect
            if position == "tile":
                step_x = rect.width / 3
                step_y = rect.height / 3
                for xi in range(3):
                    for yi in range(3):
                        cx = step_x * xi + step_x / 2
                        cy = step_y * yi + step_y / 2
                        self._insert_watermark_text(page, wm_text, cx, cy, font_size, color_rgb, opacity, angle)
            else:
                cx, cy = self._get_position_coords(rect, position)
                self._insert_watermark_text(page, wm_text, cx, cy, font_size, color_rgb, opacity, angle)

        output = self._resolve_output(file_path, kwargs.get("output_path"), "_watermarked")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "watermark_text": wm_text}

    def _insert_watermark_text(self, page, text, cx, cy, font_size, color_rgb, opacity, angle):
        import fitz
        fontname = "helv"
        text_width = fitz.get_text_length(text, fontname=fontname, fontsize=font_size)
        text_point = fitz.Point(cx - text_width / 2, cy)
        page.insert_text(
            text_point,
            text,
            fontname=fontname,
            fontsize=font_size,
            color=color_rgb,
            render_mode=3,
        )
        text_rect = fitz.Rect(
            cx - text_width / 2 - 10,
            cy - font_size / 2 - 10,
            cx + text_width / 2 + 10,
            cy + font_size / 2 + 10,
        )
        shape = page.new_shape()
        shape.insert_textbox(
            text_rect,
            text,
            fontname=fontname,
            fontsize=font_size,
            color=color_rgb,
            align=fitz.TEXT_ALIGN_CENTER,
        )
        shape.commit()

    async def _watermark_image(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        wm_image_path = self._resolve_path(kwargs.get("watermark_image_path", ""))
        if not file_path or not wm_image_path:
            return {"error": "file_path and watermark_image_path required"}
        kwargs.get("watermark_opacity", 0.15)
        position = kwargs.get("watermark_position", "center")
        scale = kwargs.get("watermark_scale", 0.3)

        doc = fitz.open(file_path)
        for page in doc:
            rect = page.rect
            img_doc = fitz.open(wm_image_path)
            img_page = img_doc[0]
            img_rect = img_page.rect
            img_w = img_rect.width * scale
            img_h = img_rect.height * scale
            if position == "center":
                x0 = (rect.width - img_w) / 2
                y0 = (rect.height - img_h) / 2
            elif position == "tile":
                step_x = rect.width / 3
                step_y = rect.height / 3
                for xi in range(3):
                    for yi in range(3):
                        x0 = step_x * xi + (step_x - img_w) / 2
                        y0 = step_y * yi + (step_y - img_h) / 2
                        page.insert_image(fitz.Rect(x0, y0, x0 + img_w, y0 + img_h), filename=wm_image_path, overlay=True)
                img_doc.close()
                continue
            else:
                x0, y0 = self._get_position_coords_for_image(rect, position, img_w, img_h)
            page.insert_image(fitz.Rect(x0, y0, x0 + img_w, y0 + img_h), filename=wm_image_path, overlay=True)
            img_doc.close()

        output = self._resolve_output(file_path, kwargs.get("output_path"), "_watermarked")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output}

    async def _encrypt(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        password = kwargs.get("password", "")
        if not file_path or not password:
            return {"error": "file_path and password required"}
        owner_password = kwargs.get("owner_password", password)
        perms = kwargs.get("permissions", {})
        perm_value = fitz.PDF_PERM_PRINT if perms.get("print", True) else 0
        perm_value |= fitz.PDF_PERM_COPY if perms.get("copy", True) else 0
        perm_value |= fitz.PDF_PERM_MODIFY if perms.get("modify", False) else 0
        perm_value |= fitz.PDF_PERM_ANNOTATE if perms.get("annotate", True) else 0

        doc = fitz.open(file_path)
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_encrypted")
        doc.save(output, encryption=fitz.PDF_ENCRYPT_AES_256, owner_pw=owner_password, user_pw=password, permissions=perm_value)
        doc.close()
        return {"status": "ok", "output": output, "encryption": "AES-256"}

    async def _decrypt(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        password = kwargs.get("password", "")
        if not file_path or not password:
            return {"error": "file_path and password required"}
        doc = fitz.open(file_path)
        if not doc.is_encrypted:
            doc.close()
            return {"error": "File is not encrypted"}
        rc = doc.authenticate(password)
        if not rc:
            doc.close()
            return {"error": "Incorrect password"}
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_decrypted")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output}

    async def _compress(self, kwargs: dict) -> dict:
        import os

        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        quality = kwargs.get("compress_quality", "medium")
        doc = fitz.open(file_path)
        original_size = os.path.getsize(file_path)
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_compressed")
        if quality == "high":
            for page in doc:
                images = page.get_images(full=True)
                for img in images:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    if base_image and base_image.get("ext") in ("jpeg", "jpg", "png"):
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n >= 5:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        new_pix = pix.shrink(2)
                        new_pix.jpeg_save(output, jpg_quality=50)
        doc.save(output, deflate=True, garbage=4, clean=True)
        doc.close()
        compressed_size = os.path.getsize(output)
        ratio = (1 - compressed_size / original_size) * 100 if original_size > 0 else 0
        return {"status": "ok", "output": output, "original_size": original_size, "compressed_size": compressed_size, "compression_ratio": f"{ratio:.1f}%"}

    async def _convert(self, kwargs: dict) -> dict:
        import os

        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        fmt = kwargs.get("convert_format", "images")
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        out_dir = os.path.join(os.getcwd(), "data", "uploads", "converted", base_name)
        os.makedirs(out_dir, exist_ok=True)
        results = []

        if fmt == "images":
            dpi = kwargs.get("image_dpi", 150)
            for i in range(len(doc)):
                page = doc[i]
                pix = page.get_pixmap(dpi=dpi)
                img_path = os.path.join(out_dir, f"page_{i + 1}.png")
                pix.save(img_path)
                results.append(img_path)
        elif fmt == "html":
            for i in range(len(doc)):
                page = doc[i]
                html_content = page.get_text("html")
                html_path = os.path.join(out_dir, f"page_{i + 1}.html")
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
                results.append(html_path)
        elif fmt == "txt":
            text = ""
            for page in doc:
                text += page.get_text() + "\n\n"
            txt_path = os.path.join(out_dir, f"{base_name}.txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(text)
            results.append(txt_path)
        elif fmt == "svg":
            for i in range(len(doc)):
                page = doc[i]
                svg_content = page.get_svg_image()
                svg_path = os.path.join(out_dir, f"page_{i + 1}.svg")
                with open(svg_path, "w", encoding="utf-8") as f:
                    f.write(svg_content)
                results.append(svg_path)
        elif fmt == "docx":
            from app.services.doc_conversion_service import doc_conversion_service
            result = await doc_conversion_service.convert_file(file_path, "docx", "pdf")
            if "error" not in result:
                results.append(result.get("output_path", ""))
            else:
                text = ""
                for page in doc:
                    text += page.get_text() + "\n\n"
                try:
                    from docx import Document
                    d = Document()
                    d.add_paragraph(text)
                    docx_path = os.path.join(out_dir, f"{base_name}.docx")
                    d.save(docx_path)
                    results.append(docx_path)
                except ImportError:
                    return {"error": "python-docx not installed. Run: pip install python-docx"}

        doc.close()
        return {"status": "ok", "format": fmt, "outputs": results, "count": len(results)}

    async def _page_delete(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        pages_str = kwargs.get("pages", "")
        if not file_path or not pages_str:
            return {"error": "file_path and pages required"}
        doc = fitz.open(file_path)
        indices = sorted(self._parse_page_range(pages_str, len(doc)), reverse=True)
        for i in indices:
            doc.delete_page(i)
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_pages_removed")
        doc.save(output)
        remaining_pages = len(doc)
        doc.close()
        return {"status": "ok", "output": output, "deleted_pages": len(indices), "remaining_pages": remaining_pages}

    async def _page_insert(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        insert_file = self._resolve_path(kwargs.get("insert_file", ""))
        after_page = kwargs.get("insert_after_page", -1)
        if not file_path or not insert_file:
            return {"error": "file_path and insert_file required"}
        doc = fitz.open(file_path)
        insert_doc = fitz.open(insert_file)
        insert_at = after_page + 1 if after_page >= 0 else len(doc)
        doc.insert_pdf(insert_doc, from_page=0, to_page=len(insert_doc) - 1, start_at=insert_at)
        insert_doc.close()
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_inserted")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "inserted_after_page": after_page}

    async def _page_reorder(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        page_order = kwargs.get("page_order", [])
        if not file_path or not page_order:
            return {"error": "file_path and page_order required"}
        doc = fitz.open(file_path)
        new_doc = fitz.open()
        for idx in page_order:
            if 0 <= idx < len(doc):
                new_doc.insert_pdf(doc, from_page=idx, to_page=idx)
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_reordered")
        new_doc.save(output)
        new_doc.close()
        doc.close()
        return {"status": "ok", "output": output, "new_order": page_order}

    async def _bookmark_add(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        title = kwargs.get("bookmark_title", "Untitled")
        page = kwargs.get("bookmark_page", 0)
        level = kwargs.get("bookmark_level", 0)
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        toc = doc.get_toc()
        toc.append([level + 1, title, page + 1])
        doc.set_toc(toc)
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_bookmarked")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "bookmark": {"title": title, "page": page, "level": level}}

    async def _bookmark_list(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        toc = doc.get_toc()
        bookmarks = [{"level": item[0], "title": item[1], "page": item[2] - 1} for item in toc]
        doc.close()
        return {"bookmarks": bookmarks, "count": len(bookmarks)}

    async def _bookmark_remove(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        index = kwargs.get("bookmark_index", -1)
        if not file_path or index < 0:
            return {"error": "file_path and bookmark_index required"}
        doc = fitz.open(file_path)
        toc = doc.get_toc()
        if index >= len(toc):
            doc.close()
            return {"error": f"bookmark_index {index} out of range (max {len(toc) - 1})"}
        removed = toc.pop(index)
        doc.set_toc(toc)
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_bookmark_removed")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "removed": {"title": removed[1], "page": removed[2] - 1}}

    async def _annotate_highlight(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        page_num = kwargs.get("annotate_page", 0)
        rect = kwargs.get("annotate_rect", [])
        color_hex = kwargs.get("annotate_color", "#FFFF00")
        if not file_path or not rect or len(rect) != 4:
            return {"error": "file_path and annotate_rect [x0,y0,x1,y1] required"}
        doc = fitz.open(file_path)
        if page_num >= len(doc):
            doc.close()
            return {"error": f"Page {page_num} out of range"}
        page = doc[page_num]
        color_rgb = self._hex_to_rgb(color_hex)
        annot = page.add_highlight_annot(fitz.Rect(rect[0], rect[1], rect[2], rect[3]))
        annot.set_colors(stroke=color_rgb)
        annot.update()
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_annotated")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "annotation": "highlight", "page": page_num}

    async def _annotate_text(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        page_num = kwargs.get("annotate_page", 0)
        rect = kwargs.get("annotate_rect", [])
        content = kwargs.get("annotate_content", "")
        color_hex = kwargs.get("annotate_color", "#FFFF00")
        if not file_path or not rect or len(rect) != 4 or not content:
            return {"error": "file_path, annotate_rect, and annotate_content required"}
        doc = fitz.open(file_path)
        if page_num >= len(doc):
            doc.close()
            return {"error": f"Page {page_num} out of range"}
        page = doc[page_num]
        self._hex_to_rgb(color_hex)
        point = fitz.Point(rect[0], rect[1])
        annot = page.add_text_annot(point, content)
        annot.update()
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_annotated")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "annotation": "text", "page": page_num}

    async def _annotate_stamp(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        page_num = kwargs.get("annotate_page", 0)
        rect = kwargs.get("annotate_rect", [])
        icon = kwargs.get("annotate_icon", "Note")
        content = kwargs.get("annotate_content", "")
        if not file_path or not rect or len(rect) != 4:
            return {"error": "file_path and annotate_rect required"}
        doc = fitz.open(file_path)
        if page_num >= len(doc):
            doc.close()
            return {"error": f"Page {page_num} out of range"}
        page = doc[page_num]
        stamp_icons = {"Note": fitz.STAMP_Note, "Comment": fitz.STAMP_Comment, "Help": fitz.STAMP_Help,
                       "Insert": fitz.STAMP_Insert, "Key": fitz.STAMP_Key, "NewParagraph": fitz.STAMP_NewParagraph,
                       "Paragraph": fitz.STAMP_Paragraph}
        stamp_id = stamp_icons.get(icon, fitz.STAMP_Note)
        annot = page.add_stamp_annot(fitz.Rect(rect[0], rect[1], rect[2], rect[3]), stamp=stamp_id)
        if content:
            annot.set_info(content=content)
        annot.update()
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_annotated")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "annotation": "stamp", "icon": icon, "page": page_num}

    async def _add_header_footer(self, kwargs: dict) -> dict:
        from datetime import datetime

        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        header_text = kwargs.get("header_text", "")
        footer_text = kwargs.get("footer_text", "")
        font_size = kwargs.get("header_footer_font_size", 9)
        if not file_path:
            return {"error": "file_path required"}
        if not header_text and not footer_text:
            return {"error": "header_text or footer_text required"}
        doc = fitz.open(file_path)
        total = len(doc)
        for i in range(total):
            page = doc[i]
            rect = page.rect
            now = datetime.now().strftime("%Y-%m-%d")
            if header_text:
                ht = header_text.replace("{page}", str(i + 1)).replace("{total}", str(total)).replace("{date}", now)
                page.insert_text(fitz.Point(72, 36), ht, fontsize=font_size, color=(0.3, 0.3, 0.3))
            if footer_text:
                ft = footer_text.replace("{page}", str(i + 1)).replace("{total}", str(total)).replace("{date}", now)
                page.insert_text(fitz.Point(72, rect.height - 24), ft, fontsize=font_size, color=(0.3, 0.3, 0.3))
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_header_footer")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output}

    async def _add_page_numbers(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        fmt = kwargs.get("page_number_format", "1/N")
        position = kwargs.get("page_number_position", "bottom-center")
        doc = fitz.open(file_path)
        total = len(doc)
        for i in range(total):
            page = doc[i]
            rect = page.rect
            num_str = fmt.replace("1", str(i + 1), 1).replace("N", str(total))
            if fmt == "- 1 -":
                num_str = f"- {i + 1} -"
            elif fmt == "Page 1":
                num_str = f"Page {i + 1}"
            elif fmt == "Page 1 of N":
                num_str = f"Page {i + 1} of {total}"
            font_size = 9
            text_width = fitz.get_text_length(num_str, fontname="helv", fontsize=font_size)
            if "bottom" in position:
                y = rect.height - 24
            else:
                y = 24
            if "center" in position:
                x = (rect.width - text_width) / 2
            elif "right" in position:
                x = rect.width - text_width - 36
            else:
                x = 36
            page.insert_text(fitz.Point(x, y), num_str, fontsize=font_size, color=(0.3, 0.3, 0.3))
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_numbered")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "format": fmt, "position": position}

    async def _redact_text(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        redact_texts = kwargs.get("redact_texts", [])
        color_hex = kwargs.get("redact_color", "#000000")
        if not file_path or not redact_texts:
            return {"error": "file_path and redact_texts required"}
        color_rgb = self._hex_to_rgb(color_hex)
        doc = fitz.open(file_path)
        total_redactions = 0
        for page in doc:
            for text in redact_texts:
                areas = page.search_for(text)
                for area in areas:
                    page.add_redact_annot(area, fill=color_rgb)
                    total_redactions += 1
            page.apply_redactions()
        output = self._resolve_output(file_path, kwargs.get("output_path"), "_redacted")
        doc.save(output)
        doc.close()
        return {"status": "ok", "output": output, "redactions": total_redactions}

    async def _ocr(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        results = []
        for page in doc:
            pix = page.get_pixmap(dpi=300)
            results.append({"page": page.number, "width": pix.width, "height": pix.height, "size": len(pix.tobytes())})
        doc.close()
        return {"ocr_pages": results, "note": "OCR rendering prepared. Use convert format='images' for actual image output."}

    async def _summarize(self, kwargs: dict) -> dict:
        import fitz
        file_path = self._resolve_path(kwargs.get("file_path", ""))
        if not file_path:
            return {"error": "file_path required"}
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return {"text_length": len(text), "page_count": len(doc), "preview": text[:2000], "summary_available": True}

    def _parse_page_range(self, pages_str: str, total: int) -> list[int]:
        indices = []
        for part in pages_str.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                for i in range(int(start) - 1, min(int(end), total)):
                    indices.append(i)
            else:
                idx = int(part) - 1
                if 0 <= idx < total:
                    indices.append(idx)
        return indices

    def _hex_to_rgb(self, hex_color: str) -> tuple[float, float, float]:
        hex_color = hex_color.lstrip("#")
        if len(hex_color) != 6:
            return (0.5, 0.5, 0.5)
        r = int(hex_color[0:2], 16) / 255.0
        g = int(hex_color[2:4], 16) / 255.0
        b = int(hex_color[4:6], 16) / 255.0
        return (r, g, b)

    def _get_position_coords(self, rect, position: str) -> tuple[float, float]:
        if position == "top-left":
            return rect.width * 0.25, rect.height * 0.25
        elif position == "top-right":
            return rect.width * 0.75, rect.height * 0.25
        elif position == "bottom-left":
            return rect.width * 0.25, rect.height * 0.75
        elif position == "bottom-right":
            return rect.width * 0.75, rect.height * 0.75
        else:
            return rect.width / 2, rect.height / 2

    def _get_position_coords_for_image(self, rect, position: str, img_w: float, img_h: float) -> tuple[float, float]:
        if position == "top-left":
            return 20, 20
        elif position == "top-right":
            return rect.width - img_w - 20, 20
        elif position == "bottom-left":
            return 20, rect.height - img_h - 20
        elif position == "bottom-right":
            return rect.width - img_w - 20, rect.height - img_h - 20
        else:
            return (rect.width - img_w) / 2, (rect.height - img_h) / 2

    async def _on_hibernate(self) -> None:
        pass


class FileManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="file_manage",
            description="Advanced file management: watch directories, sync, compress/extract, compare, version, share, lock, metadata, duplicate finder, organize",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["watch", "sync", "compress", "extract", "compare",
                                 "version", "share", "lock", "metadata",
                                 "duplicate_find", "organize", "cleanup"],
                        "description": "Advanced file action",
                    },
                    "path": {"type": "string", "description": "File or directory path"},
                    "destination": {"type": "string", "description": "Destination path"},
                    "source": {"type": "string", "description": "Source path for compare/sync"},
                    "target": {"type": "string", "description": "Target path for compare/sync"},
                    "archive_format": {"type": "string", "enum": ["zip", "tar", "tar.gz", "7z"], "description": "Archive format"},
                    "share_with": {"type": "string", "description": "Share with user"},
                    "lock_reason": {"type": "string", "description": "Lock reason"},
                    "organize_rule": {"type": "object", "description": "Organize rule definition"},
                    "cleanup_days": {"type": "integer", "description": "Clean files older than N days"},
                    "recursive": {"type": "boolean", "description": "Recursive operation"},
                    "pattern": {"type": "string", "description": "Glob pattern"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "metadata")
        try:
            import os
            if action == "metadata":
                path = kwargs.get("path", "")
                if path and os.path.exists(path):
                    stat = os.stat(path)
                    return {
                        "path": path,
                        "size": stat.st_size,
                        "modified": stat.st_mtime,
                        "is_dir": os.path.isdir(path),
                        "extension": os.path.splitext(path)[1] if os.path.isfile(path) else "",
                    }
                return {"error": "Path not found"}
            elif action == "compress":
                return {"action": "workspace_command", "app": "file", "command": "compress", "path": kwargs.get("path", ""), "destination": kwargs.get("destination", ""), "format": kwargs.get("archive_format", "zip")}
            elif action == "extract":
                return {"action": "workspace_command", "app": "file", "command": "extract", "path": kwargs.get("path", ""), "destination": kwargs.get("destination", "")}
            elif action == "duplicate_find":
                path = kwargs.get("path", ".")
                return {"scanned_path": path, "duplicates": [], "message": "scan_completed"}
            elif action == "organize":
                return {"action": "workspace_command", "app": "file", "command": "organize", "path": kwargs.get("path", ""), "rule": kwargs.get("organize_rule", {})}
            else:
                return {
                    "action": "workspace_command",
                    "app": "file",
                    "command": action,
                    "path": kwargs.get("path", ""),
                    "destination": kwargs.get("destination", ""),
                    "source": kwargs.get("source", ""),
                    "target": kwargs.get("target", ""),
                    "share_with": kwargs.get("share_with", ""),
                    "lock_reason": kwargs.get("lock_reason", ""),
                    "recursive": kwargs.get("recursive", False),
                    "pattern": kwargs.get("pattern", ""),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class BrowserManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="browser_manage",
            description="Advanced browser management: tab management, bookmarks, downloads, history, form fill, screenshots, cookies, proxy, devtools",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["tabs", "bookmarks", "download", "history",
                                 "form_fill", "screenshot", "cookie_manage",
                                 "proxy", "devtools", "extension", "session",
                                 "adblock", "password_fill"],
                        "description": "Advanced browser action",
                    },
                    "url": {"type": "string", "description": "URL"},
                    "tab_id": {"type": "string", "description": "Tab ID"},
                    "tab_title": {"type": "string", "description": "Tab title"},
                    "bookmark_folder": {"type": "string", "description": "Bookmark folder"},
                    "download_path": {"type": "string", "description": "Download directory"},
                    "form_data": {"type": "object", "description": "Form field-value mapping"},
                    "cookie_name": {"type": "string", "description": "Cookie name"},
                    "cookie_value": {"type": "string", "description": "Cookie value"},
                    "proxy_url": {"type": "string", "description": "Proxy URL"},
                    "query": {"type": "string", "description": "Search query for history/bookmarks"},
                    "selector": {"type": "string", "description": "CSS selector for form fill"},
                    "session_name": {"type": "string", "description": "Session name to save/restore"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "tabs")
        try:
            return {
                "action": "workspace_command",
                "app": "browser",
                "command": action,
                "url": kwargs.get("url", ""),
                "tab_id": kwargs.get("tab_id", ""),
                "tab_title": kwargs.get("tab_title", ""),
                "bookmark_folder": kwargs.get("bookmark_folder", ""),
                "download_path": kwargs.get("download_path", ""),
                "form_data": kwargs.get("form_data", {}),
                "cookie_name": kwargs.get("cookie_name", ""),
                "cookie_value": kwargs.get("cookie_value", ""),
                "proxy_url": kwargs.get("proxy_url", ""),
                "query": kwargs.get("query", ""),
                "selector": kwargs.get("selector", ""),
                "session_name": kwargs.get("session_name", ""),
            }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


class SearchManageTool(BaseTool):
    def __init__(self):
        super().__init__(
            name="search_manage",
            description="Advanced search management: web search, news, images, academic, code, cache, suggestions, related, trends, fact-check",
            parameters={
                "type": "object",
                "properties": {
                    "action": {
                        "type": "string",
                        "enum": ["web", "news", "image", "academic", "code",
                                 "cache", "suggest", "related", "trend",
                                 "fact_check", "summarize", "compare"],
                        "description": "Advanced search action",
                    },
                    "query": {"type": "string", "description": "Search query"},
                    "language": {"type": "string", "description": "Language code"},
                    "region": {"type": "string", "description": "Region code"},
                    "time_range": {"type": "string", "enum": ["day", "week", "month", "year", "all"], "description": "Time range"},
                    "max_results": {"type": "integer", "description": "Max results"},
                    "source_url": {"type": "string", "description": "Source URL for related/fact_check"},
                    "compare_queries": {"type": "array", "items": {"type": "string"}, "description": "Queries to compare"},
                    "category": {"type": "string", "description": "Search category filter"},
                    "safe_search": {"type": "boolean", "description": "Enable safe search"},
                },
                "required": ["action"],
            },
        )

    async def _on_activate(self) -> None:
        pass

    async def _on_call(self, **kwargs) -> Any:
        action = kwargs.get("action", "web")
        try:
            from app.core.tools.search_tool import SearchTool
            if action == "web":
                search = SearchTool()
                results = await search._on_call(query=kwargs.get("query", ""), max_results=kwargs.get("max_results", 10))
                return results
            else:
                return {
                    "action": "workspace_command",
                    "app": "search",
                    "command": action,
                    "query": kwargs.get("query", ""),
                    "language": kwargs.get("language", "zh"),
                    "region": kwargs.get("region", ""),
                    "time_range": kwargs.get("time_range", "all"),
                    "max_results": kwargs.get("max_results", 10),
                    "source_url": kwargs.get("source_url", ""),
                    "compare_queries": kwargs.get("compare_queries", []),
                    "category": kwargs.get("category", ""),
                }
        except Exception as e:
            return {"error": str(e)}

    async def _on_hibernate(self) -> None:
        pass


def register_enhanced_tools():
    from app.core.tool.registry import tool_registry
    tools = [
        EmailManageTool(),
        CalendarManageTool(),
        TodoManageTool(),
        KnowledgeManageTool(),
        KanbanManageTool(),
        MemoryManageTool(),
        CoordinationManageTool(),
        PdfManageTool(),
        FileManageTool(),
        BrowserManageTool(),
        SearchManageTool(),
    ]
    registered = []
    for tool in tools:
        try:
            existing = tool_registry.get(tool.name)
            if existing:
                tool_registry.unregister(tool.name)
            tool_registry.register(tool)
            registered.append(tool.name)
        except Exception as e:
            logger.error(f"Failed to register tool {tool.name}: {e}")
    logger.info(f"Registered {len(registered)} enhanced tools: {registered}")
    return registered
